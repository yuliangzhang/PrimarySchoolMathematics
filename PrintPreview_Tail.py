#!/usr/bin/env python3
# -*- coding=utf-8 -*-

# @Time    : 2018-11-02
# @Author  : J.sky
# @Mail    : bosichong@qq.com
# @Site    : www.2vv.net
# @Title   : 基于Python开发的小学生口算题生成器
# @Url     : http://2vv.net/blog/83.html
# @Details : Python实现小学生加减乘除速算考试题卷。
# @Other   : OS X 10.11.6
#            Python 3.6.1
#            vscode


'''
孩子上小学一年级了，加减乘除的口算就要开始练习了，估计老题肯定会让家长出题，所以提前准备一下.

利用Python开发了一套自动生成小学生口算题的小应用。而且今天是程序员节，撸200行代码庆祝一下。：）

程序核心功能：

    1.根据条件生成相关的口算题.

    2.保存为.docx用来打印.


开心Python Django 学习交流q群：217840699


Author  : J.sky
Mail    : bosichong@qq.com


'''
import os

from docx import Document  # 引入docx类生成docx文档
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

__version__ = "1.0.1"


class PrintPreview:
    '''本类负责生成完整的口算题文档使之适合打印机打印。可以生成多套题，生成数可以控。

    - @p_list   list
    需要打印口算题库，至少包含一套口算题

    - @p_title   list
    页面标题，这个标题的生成依据程序题型的选择和数字的范围选择而生成，例如：选择了0-20，加减法，进退位
    则自动生成标题为：0到20加减法进退位混合口算题，list中包含了多套题的页面标题名称

    - @p_column  int
    打印页排版口算题的列数

    '''


    def __init__(self, l, tit, subtitle, col=3, tsize=20, subsize=11, csize=16,
                # 默认输出文件地址为项目根目录 
                docxpath=os.path.join(os.path.dirname(os.path.abspath(__file__)),"")):
        '''
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题
        :param subtitle str 小标题
        :param col: int 列数
        :param tsize: int 标题字号
        :param csize: int 口算题字号
        :param docxpath str 保存路径
        
        '''
        self.p_list = l
        self.p_title = tit
        self.p_subtitle = subtitle
        self.p_column = col
        self.p_title_size = tsize
        self.p_subtitle_size = subsize
        self.p_content_siae = csize
        self.docxpath = docxpath

    def create_psmdocx(self, l, title, docxname):
        '''
        :param l list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        '''
        # if (title == ''):
        #     page_title = '小学生寒假作业打卡题'
        # else:
        #     page_title = title

        kousuan = l[0]
        compare = l[1]
        application = l[2]

        page_title = title
        p_docx = Document()  # 创建一个docx文档
        p_docx.styles['Normal'].font.name = u'宋体'  # 可换成word里面任意字体
        p_docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        p = p_docx.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run(page_title)
        run.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应
        run.bold = True
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


        sp = p_docx.add_paragraph()
        sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        srun = sp.add_run(self.p_subtitle)
        srun.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun.font.size = Pt(self.p_subtitle_size)  # 字体大小设置，和word里面的字号相对应
        srun.font.name = u'宋体'
        srun._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


        '''
            口算题部分
        '''
        sp1 = p_docx.add_paragraph()
        sp1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
        srun1 = sp1.add_run("一、口算题")
        srun1.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun1.font.size = Pt(16)  # 字体大小设置，和word里面的字号相对应
        srun1.bold=True
        srun1.font.name = u'宋体'
        srun1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


        # 判断需要用到的行数
        if (len(kousuan) % self.p_column):
            rs = len(kousuan) // self.p_column + 2
        else:
            rs = len(kousuan) // self.p_column +1

        # print(rs)

        # 将口算题添加到docx表格中
        k = 0  # 计数器
        table1 = p_docx.add_table(rows=rs, cols=self.p_column)

        for i in range(rs):
            if i >= 0:
                row_cells = table1.rows[i].cells
                for j in range(self.p_column):
                    if (k > len(kousuan) - 1):
                        break
                    else:
                        row_cells[j].text = kousuan[k]
                        k = k + 1
        table1.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table1.style.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        table1.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应


        '''
            比较大小的题目内容
        '''
        sp2 = p_docx.add_paragraph()
        sp2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
        srun2 = sp2.add_run("二、比较大小")
        srun2.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun2.font.size = Pt(16)  # 字体大小设置，和word里面的字号相对应
        srun2.bold = True
        srun2.font.name = u'宋体'
        srun2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # 判断需要用到的行数
        if (len(compare) % self.p_column):
            rs = len(compare) // self.p_column + 2
        else:
            rs = len(compare) // self.p_column + 1

        # print(rs)

        # 将口算题添加到docx表格中
        k = 0  # 计数器
        table2 = p_docx.add_table(rows=rs, cols=self.p_column)

        for i in range(rs):
            if i >= 0:
                row_cells = table2.rows[i].cells
                for j in range(self.p_column):
                    if (k > len(compare) - 1):
                        break
                    else:
                        row_cells[j].text = compare[k]
                        k = k + 1
        table2.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table2.style.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        table2.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应

        '''
            应用题部分
        '''

        sp3 = p_docx.add_paragraph()
        sp3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
        srun3 = sp3.add_run("三、应用题")
        srun3.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun3.font.size = Pt(16)  # 字体大小设置，和word里面的字号相对应
        srun3.bold=True
        srun3.font.name = u'宋体'
        srun3._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        for app in application:
            sp4 = p_docx.add_paragraph()
            sp4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
            srun4 = sp4.add_run(app + '\n')
            srun4.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
            srun4.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
            srun3.font.name = u'宋体'
            srun3._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # sp5 = p_docx.add_paragraph()
            # srun5 = sp5.add_run('\n')
            # srun5.font.color.rgb = RGBColor(20, 0, 0)  # 颜色设置，这里是用RGB颜色
            # srun5.font.size = Pt(12)

        print(self.docxpath+'{}.docx'.format(docxname))
        p_docx.save(self.docxpath+'{}.docx'.format(docxname))  # 输出docx

    def produce(self):
        k = 1
        for l, t in zip(self.p_list, self.p_title):
            self.create_psmdocx(l, t, t + str(k))
    def produce_docx(self):
        self.create_psmdocx(self.p_list, self.p_title[0], self.p_title[0])




if __name__ == '__main__':
    l = [['1-17=', '3-4=', '13-6=', '15-5=', '2-4=', '15-9=', '12-13=', '15-12=', '14-16=', '4-11=', '18-16=', '12-14=',
          ],
         ['1-17=', '3-4=', '13-6=', '15-5=', '2-4=', '15-9=', '12-13=', '15-12=', '14-16=', '4-11=', '18-16=', '12-14=',
          '14-7=', '7-17=', '16-19=',  ]]
    t = ['小学生口算题', '小学生口算题']
    pp = PrintPreview(l, t,"姓名：__________ 日期：____月____日 时间：________ 对题：____道",)
    pp.produce()
